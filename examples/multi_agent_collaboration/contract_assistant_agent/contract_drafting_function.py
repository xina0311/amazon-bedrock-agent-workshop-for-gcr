import json
from datetime import datetime, timedelta
import boto3
import uuid
import os
from io import StringIO
import re
from io import BytesIO
from docx import Document

bucket_name = os.environ["BUCKET_NAME"]


def get_template_from_s3(contract_type):
    """
    Retrieve the contract template from S3 and return the Document object

    Args:
        contract_type (str): Type of contract (e.g., 'purchase', 'service', etc.)

    Returns:
        docx.Document: Document object if successful, None if failed
    """
    s3 = boto3.client("s3")
    template_key = f"contract_templates/{contract_type}_agreement.docx"

    try:
        # Get the template file from S3
        response = s3.get_object(Bucket=bucket_name, Key=template_key)

        # Read the binary content
        doc_binary = response["Body"].read()

        # Create a BytesIO object
        doc_file = BytesIO(doc_binary)

        # Return the Document object
        return Document(doc_file)

    except Exception as e:
        print(f"Error reading template from S3: {str(e)}")
        return None


def draft_contract(contract_type, contract_details):
    try:
        # Get the base template content from S3
        print(f"Contract type: {contract_type}")
        template_content = get_template_from_s3(contract_type)

        if not template_content:
            return {"error": f"Template not found for contract type: {contract_type}"}

        # Prepare prompt for the LLM
        prompt = f"""
        Act as a legal contract expert. Using this contract template:
        {template_content}
        
        And these contract details:
        {contract_details}
        
        Generate a professional contract that:
        1. Maintains legal compliance
        2. Uses appropriate legal terminology
        3. Expands clauses where necessary
        4. Adds relevant standard clauses
        5. Ensures all terms are clearly defined
        
        Return only the generated contract text without any additional explanations.
        """

        # Call Bedrock using converse API
        bedrock_client = bedrock_client = boto3.client(service_name="bedrock-runtime")
        model_id = "anthropic.claude-3-haiku-20240307-v1:0"
        response = bedrock_client.converse(
            modelId=model_id,
            messages=[{"role": "user", "content": [{"text": prompt}]}],
            inferenceConfig={"temperature": 0, "maxTokens": 1024, "topP": 0.9},
            # additionalModelRequestFields={
            # }
        )

        # Extract the generated content
        try:
            generated_content = response["output"]["message"]["content"][0]["text"]
        except Exception as e:
            print(f"Error parsing model response: {str(e)}")
            return {"error": "Failed to parse model response"}

        # Create contract object with metadata
        contract = {
            "contract_id": str(uuid.uuid4()),
            "contract_type": contract_type,
            "created_date": datetime.now().strftime("%B %d, %Y"),
            "content": generated_content,
            "details": contract_details,
            "status": "DRAFT",
            "model_metrics": {
                "latency_ms": response.get("metrics", {}).get("latencyMs"),
                "input_tokens": response.get("usage", {}).get("inputTokens"),
                "output_tokens": response.get("usage", {}).get("outputTokens"),
            },
        }

        return contract

    except Exception as e:
        print(f"Error in draft_contract: {str(e)}")
        return {"error": f"Error creating contract: {str(e)}"}

# draft_contract 1 is an option of drafting the contract without using LLM. The method works but if the payload is larger than the current lambda limit of 25kb, it throws exception. I just put in case customer want to try this with quota increase.
def draft_contract1(contract_type, contract_details):
    """
    Draft a contract using template and user provided details

    Args:
        contract_type (str): Type of contract (e.g., 'purchase', 'franchise', 'timeandmaterial')
        contract_details (dict): Contract specific details to populate template

    Returns:
        dict: Contract object with content and metadata
    """
    try:
        # Get template document
        template_doc = get_template_from_s3(contract_type)
        if not template_doc:
            raise Exception(f"Could not load template for {contract_type}")

        # Create new document for the contract
        doc = Document()

        # Prepare template variables
        template_vars = {
            "contract_id": str(uuid.uuid4()),
            "created_date": datetime.now().strftime("%B %d, %Y"),
            **contract_details,
        }

        # Process template paragraphs and replace placeholders
        for paragraph in template_doc.paragraphs:
            new_para = doc.add_paragraph()
            # Copy paragraph format
            new_para.style = paragraph.style
            new_para.alignment = paragraph.alignment

            text = paragraph.text
            # Replace placeholders in format [VARIABLE_NAME] or {variable_name}
            for key, value in template_vars.items():
                # Handle both [KEY] and {key} formats
                text = text.replace(f"[{key.upper()}]", str(value))
                text = text.replace(f"{{{key}}}", str(value))

            # Add replaced text while maintaining runs (formatting)
            if text.strip():
                new_para.add_run(text)

        # Process template tables
        for table in template_doc.tables:
            new_table = doc.add_table(rows=len(table.rows), cols=len(table.columns))
            new_table.style = table.style

            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    text = cell.text
                    # Replace placeholders in tables
                    for key, value in template_vars.items():
                        text = text.replace(f"[{key.upper()}]", str(value))
                        text = text.replace(f"{{{key}}}", str(value))
                    new_table.cell(i, j).text = text

        # Convert document to bytes
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        # Create contract object
        contract = {
            "contract_id": template_vars["contract_id"],
            "contract_type": contract_type,
            "created_date": template_vars["created_date"],
            "content": doc_bytes.getvalue(),
            "details": contract_details,
            "status": "DRAFT",
        }

        return contract

    except Exception as e:
        print(f"Error drafting contract: {str(e)}")
        return None


# Helper function to format currency values
def format_currency(amount):
    try:
        return "${:,.2f}".format(float(str(amount).replace(",", "")))
    except:
        return amount


# Helper function to validate and format dates
def format_date(date_str):
    try:
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")
        return date_obj.strftime("%B %d, %Y")
    except:
        return date_str


def get_specific_template_fields(contract_type):
    """
    Get required fields for specific contract type
    """
    fields = {
        "purchase": [
            "buyer_name",
            "seller_name",
            "property_address",
            "purchase_price",
            "closing_date",
            "deposit_amount",
        ],
        "franchise": [
            "franchisor_name",
            "franchisee_name",
            "business_type",
            "brand_name",
            "territory",
            "initial_fee",
            "royalty_rate",
        ],
        "timeandmaterial": [
            "contractor_name",
            "client_name",
            "scope_of_work",
            "labor_rates",
            "markup_percentage",
            "not_to_exceed_amount",
        ],
    }
    return fields.get(contract_type, [])


def save_contract_to_s3(contract, bucket_name=bucket_name):
    """
    Save contract as a DOCX file to S3

    Args:
        contract (dict): Contract dictionary containing:
            - contract_id (str): Unique identifier for the contract
            - content (dict): Contract content with sections/paragraphs
            - metadata (dict, optional): Additional contract metadata
        bucket_name (str): S3 bucket name

    Returns:
        dict: Status of the upload operation
    """
    s3 = boto3.client("s3")
    contract_id = contract["contract_id"]
    file_name = f"contracts/{contract_id}.docx"

    try:
        # Create a new Document
        doc = Document()

        # Add title
        doc.add_heading(f"Contract #{contract_id}", 0)

        # Add contract content
        if "metadata" in contract:
            doc.add_heading("Contract Details", level=1)
            for key, value in contract["metadata"].items():
                doc.add_paragraph(f"{key}: {value}")

            # Add separator
            doc.add_paragraph()

        # Add main content
        if isinstance(contract["content"], str):
            # If content is a simple string
            doc.add_paragraph(contract["content"])
        else:
            # If content is structured (e.g., sections, paragraphs)
            for section in contract["content"].get("sections", []):
                if "title" in section:
                    doc.add_heading(section["title"], level=1)
                if "content" in section:
                    doc.add_paragraph(section["content"])

            # Add tables if present
            for table_data in contract["content"].get("tables", []):
                if table_data:
                    rows = len(table_data)
                    cols = len(table_data[0]) if table_data else 0
                    if rows and cols:
                        table = doc.add_table(rows=rows, cols=cols)
                        for i, row_data in enumerate(table_data):
                            for j, cell_data in enumerate(row_data):
                                table.cell(i, j).text = str(cell_data)

        # Convert document to bytes
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        # Upload to S3
        s3.put_object(
            Bucket=bucket_name,
            Key=file_name,
            Body=doc_bytes.getvalue(),
            ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        return {
            "status": "success",
            "location": f"s3://{bucket_name}/{file_name}",
            "file_type": "docx",
        }

    except Exception as e:
        return {"status": "error", "message": str(e)}


def get_named_parameter(event, name):
    if "parameters" in event:
        return next(
            (item["value"] for item in event["parameters"] if item["name"] == name),
            None,
        )
    return None


def populate_function_response(event, response_body):
    return {
        "response": {
            "actionGroup": event["actionGroup"],
            "function": event["function"],
            "functionResponse": {
                "responseBody": {"TEXT": {"body": str(response_body)}}
            },
        }
    }


def determine_contract_type(requirements):
    # Analyze requirements to determine appropriate contract type
    requirements = requirements.lower()
    if any(term in requirements for term in ["purchase", "buy", "goods", "product"]):
        return "purchase"
    elif any(term in requirements for term in ["franchise", "brand", "license"]):
        return "franchise"
    elif any(
        term in requirements
        for term in ["consulting", "hourly", "time", "material", "service"]
    ):
        return "timeandmaterial"
    return "unknown"


def get_contract_questions(contract_type):
    questions = {
        "purchase": [
            "What is the buyer entity name?",
            "What is the seller entity name?",
            "What are the goods/products being purchased?",
            "What is the total purchase amount?",
            "What are the delivery terms?",
            "What are the payment terms?",
        ],
        "franchise": [
            "What is the franchisor name?",
            "What is the franchisee name?",
            "What is the franchise territory?",
            "What is the initial franchise fee?",
            "What is the royalty percentage?",
            "What is the term length in years?",
        ],
        "timeandmaterial": [
            "What is the service provider name?",
            "What is the client name?",
            "What is the hourly rate?",
            "What is the estimated number of hours?",
            "What is the payment schedule?",
            "What is the scope of work?",
        ],
    }
    return questions.get(contract_type, [])


def lambda_handler(event, context):
    print(event)
    function = event["function"]

    if function == "determine_contract_type":
        requirements = get_named_parameter(event, "requirements")
        if not requirements:
            result = "Please provide requirements to determine contract type"
        else:
            result = determine_contract_type(requirements)

    elif function == "get_contract_questions":
        contract_type = get_named_parameter(event, "contract_type")
        if not contract_type:
            result = "Please specify contract type"
        else:
            result = get_contract_questions(contract_type)

    elif function == "draft_contract":
        contract_type = get_named_parameter(event, "contract_type")
        contract_details = get_named_parameter(event, "contract_details")

        if not contract_type or not contract_details:
            result = "Please provide both contract type and details"
        else:
            try:
                contract_details = json.loads(contract_details)
                contract = draft_contract(contract_type, contract_details)
                save_result = save_contract_to_s3(contract)
                result = {"contract": contract, "save_status": save_result}
            except json.JSONDecodeError:
                result = "Invalid contract details format"
            except Exception as e:
                result = f"Error creating contract: {str(e)}"

    else:
        result = f"Unrecognized function: {function}"

    response = populate_function_response(event, result)
    print(response)
    return response
